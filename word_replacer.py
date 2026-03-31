from __future__ import annotations

from dataclasses import dataclass
import re
from typing import Iterable, List



RULE_PATTERN = re.compile(r"^(literal|regex)\s*:\s*(.+?)\s*=>\s*(.*)$")


@dataclass(frozen=True)
class ReplacementRule:
    mode: str
    source: str
    target: str

    def apply(self, text: str) -> str:
        if self.mode == "literal":
            return text.replace(self.source, self.target)
        if self.mode == "regex":
            normalized_target = re.sub(r"\$(\d+)", r"\\g<\1>", self.target)
            return re.sub(self.source, normalized_target, text)
        raise ValueError(f"Unsupported mode: {self.mode}")


def parse_rules(raw_rules: str) -> List[ReplacementRule]:
    """
    Parse replacement rules.

    Format (one rule per line):
      literal:旧文本=>新文本
      regex:正则表达式=>新文本

    Lines starting with # and blank lines are ignored.
    """
    rules: List[ReplacementRule] = []
    for line_no, line in enumerate(raw_rules.splitlines(), start=1):
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue

        match = RULE_PATTERN.match(stripped)
        if not match:
            raise ValueError(
                f"第 {line_no} 行格式错误。请使用 'literal:旧=>新' 或 'regex:模式=>新'。"
            )

        mode, source, target = match.groups()
        if mode == "literal" and not source:
            raise ValueError(f"第 {line_no} 行 literal 规则的旧文本不能为空。")
        if mode == "regex":
            try:
                re.compile(source)
            except re.error as exc:
                raise ValueError(f"第 {line_no} 行正则表达式无效: {exc}") from exc

        rules.append(ReplacementRule(mode=mode, source=source, target=target))

    if not rules:
        raise ValueError("未检测到可用规则。请至少提供一条规则。")

    return rules


def _iter_text_blocks(document) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def apply_rules_to_docx(input_path: str, output_path: str, rules: List[ReplacementRule]) -> int:
    from docx import Document

    doc = Document(input_path)
    total_changes = 0

    for paragraph in _iter_text_blocks(doc):
        original = paragraph.text
        updated = original
        for rule in rules:
            updated = rule.apply(updated)

        if updated != original:
            total_changes += 1
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = updated
            else:
                paragraph.add_run(updated)

    doc.save(output_path)
    return total_changes
